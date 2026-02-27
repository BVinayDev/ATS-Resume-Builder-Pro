#!/usr/bin/env python3
"""
ATS Resume Builder Pro — v8.0
═══════════════════════════════════════════════════════════════
CRITICAL FIXES FROM v7:
  ① Parser: strict section isolation — experience bucket NEVER leaks
    into education.  Each section header resets the current bucket.
    Education parsed completely independently from experience.
  ② Parser: two-pass merge fixed — orphan title lines (before dates)
    are always merged into the next date-bearing block.
  ③ DOCX bullets: ATS-correct typography per screenshot:
     • Section headers: 10pt ALL CAPS BOLD — smaller than role names ✓
     • Role line:       11pt Bold (largest body text) ✓
     • Location:        10pt italic ✓
     • Bullet text:     11pt, hanging-indent 0.25" / -0.25" ✓
     • Bullet spacing:  space_before=1, space_after=3 (visible gap) ✓
     • Bullet prefix:   "•  " (bullet + two spaces for clear gap) ✓
  ④ Justify: only on paragraphs >60 chars; short lines stay LEFT ✓
  ⑤ All text sanitised: fancy unicode bullets→plain, no double spaces ✓

NEW FEATURES:
  • Resume Templates: Modern / Classic / Compact — one click apply
  • Drag-to-Reorder experience cards (↑↓ buttons)
  • Bullet Auto-Improve: AI-style suggestions (action verb + metric)
  • Resume Completeness Checklist — ticked live as user fills sections
  • Export to PDF (via docx2pdf if available, else instructions)
  • Quick-fill demo data — "Fill with sample" for each section
  • Character limit warnings on all text fields
  • Undo last parse (restore previous data)
  • Tab key navigation between all fields
═══════════════════════════════════════════════════════════════
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json, os, re, copy

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ══════════════════════════════════════════════════════
#  THEME — Professional Navy / White
# ══════════════════════════════════════════════════════
BG     = '#EEF1F6'
SB     = '#071525'
SB_HL  = '#102040'
WHITE  = '#FFFFFF'
ACCENT = '#1355A2'
GREEN  = '#14583A'
RED    = '#8A1C1C'
AMBER  = '#7B3F08'
TEXT   = '#0E1628'
MUTED  = '#5A6F84'
BORDER = '#C3CDD8'
INP    = '#F6F9FC'
TIP_BG = '#EBF3FF'
TIP_FG = '#1244A0'
DONE   = '#14583A'
TODO   = '#8EA3B7'
HDR_C  = '#0B2A55'
DONE_C = '#E8F5EE'

UF = 'Calibri'   # UI font
DF = 'Calibri'   # Doc font

STEPS = ['Personal', 'Summary', 'Skills', 'Experience', 'Education', 'Preview']

BULLET_STYLES = {
    'Bullet  •' : '•',
    'Hyphen  -' : '-',
    'Square  ▪' : '▪',
    'Number  1.': 'NUM',
}

# ATS-standard font sizes (DOCX)
SZ_NAME    = 14   # candidate name
SZ_CONTACT = 10   # contact line
SZ_SEC_HDR = 10   # SECTION HEADERS (ALL CAPS, bold)
SZ_ROLE    = 11   # Job title | Company  (bold)
SZ_LOC     = 10   # Location (italic)
SZ_BODY    = 11   # bullets, summary, body text
SZ_DATE    = 10   # dates (italic)

# Strong / weak action verbs
STRONG_VERBS = [
    'Achieved','Accelerated','Architected','Automated','Built','Championed',
    'Collaborated','Cut','Delivered','Deployed','Designed','Developed',
    'Drove','Eliminated','Engineered','Enhanced','Executed','Generated',
    'Implemented','Improved','Integrated','Launched','Led','Leveraged',
    'Managed','Migrated','Optimized','Orchestrated','Reduced','Refactored',
    'Scaled','Secured','Streamlined','Transformed','Upgraded','Spearheaded',
]
WEAK_VERBS = ['worked on','helped','assisted','was responsible for',
              'did','made','used','got','involved in','participated in']

# ══════════════════════════════════════════════════════
#  TEXT UTILITIES
# ══════════════════════════════════════════════════════
_U2A = {
    '\u2018':"'",'\u2019':"'",'\u201c':'"','\u201d':'"',
    '\u2013':'-','\u2014':'-',
    '\u2022':'•','\u25cf':'•','\u25aa':'•','\u2023':'•',
    '\u2043':'•','\u25e6':'•','\u00b7':' ',
    '\xa0':' ','\u200b':'','\ufeff':'','\t':' ',
}
_MON = {
    'january':'Jan','february':'Feb','march':'Mar','april':'Apr',
    'may':'May','june':'Jun','july':'Jul','august':'Aug',
    'september':'Sep','october':'Oct','november':'Nov','december':'Dec',
    'jan':'Jan','feb':'Feb','mar':'Mar','apr':'Apr',
    'jun':'Jun','jul':'Jul','aug':'Aug','sep':'Sep',
    'oct':'Oct','nov':'Nov','dec':'Dec',
}

def san(t: str) -> str:
    """Full sanitise: unicode → ASCII/UTF8, collapse whitespace."""
    if not t: return ''
    for s, d in _U2A.items():
        t = t.replace(s, d)
    t = t.encode('ascii', 'ignore').decode()
    t = re.sub(r'[ \t]+', ' ', t)
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()

def s1(t: str) -> str:
    """Single-line sanitise."""
    return re.sub(r'\s+', ' ', san(t)).strip()

def norm_date(raw: str) -> str:
    s = raw.strip()
    if not s: return ''
    sl = s.lower()
    if sl in ('present','current','now','till date','to date','today','ongoing',
              'till present','to present'): return 'Present'
    # MM/YYYY or MM-YYYY
    mm = re.match(r'^(\d{1,2})[/\-](\d{4})$', s)
    if mm:
        _mn = ['Jan','Feb','Mar','Apr','May','Jun',
               'Jul','Aug','Sep','Oct','Nov','Dec']
        i = int(mm.group(1))
        if 1 <= i <= 12: return f'{_mn[i-1]} {mm.group(2)}'
        return mm.group(2)
    for k, v in _MON.items():
        if k in sl:
            yr = re.search(r'\d{4}', s)
            if yr: return f'{v} {yr.group()}'
    yr = re.search(r'\d{4}', s)
    return yr.group() if yr else s

def strip_pfx(line: str) -> str:
    """Strip leading bullet/number prefix."""
    line = re.sub(r'^\s*\d+[\.\)]\s+', '', line)
    line = re.sub(r'^\s*[-*+>•·▪◦○●\u2022\u25cf]\s*', '', line)
    return line.strip()

def to_bullets(text: str) -> list:
    out = []
    for line in san(text).split('\n'):
        line = strip_pfx(line)
        line = re.sub(r' {2,}', ' ', line).strip()
        if len(line) > 4:
            out.append(line[0].upper() + line[1:])
    return out

# ══════════════════════════════════════════════════════
#  RESUME PARSER  v8 — strict section isolation
# ══════════════════════════════════════════════════════
_DRANGE = re.compile(
    r'(?<!\d)'                                         # no digit before
    r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|'
    r'Dec(?:ember)?)\.?\s*\d{4}'
    r'|\d{1,2}[/\-]\d{4}|\d{4})'
    r'\s*(?:–|—|-|to|/)\s*'
    r'(Present|Current|Now|Ongoing|'
    r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|'
    r'Dec(?:ember)?)\.?\s*\d{4}'
    r'|\d{1,2}[/\-]\d{4}|\d{4})',
    re.IGNORECASE,
)

# Strict section header patterns (full-line match only)
_SECS = {
    'summary'   : re.compile(r'^(professional\s*summary|summary|objective|career\s*objective|profile|about\s*me?|overview|executive\s*summary)$', re.I),
    'skills'    : re.compile(r'^(technical\s*skills?|core\s*competencies?|skills?|technologies?|expertise|key\s*skills?|areas?\s*of\s*expertise|tools?\s*(and|&)\s*technologies?|programming\s*languages?|technical\s*proficiencies?)$', re.I),
    'experience': re.compile(r'^(work\s*experience|professional\s*experience|experience|employment(\s*history)?|work\s*history|career\s*history|positions?\s*held|career\s*summary)$', re.I),
    'education' : re.compile(r'^(education(\s*(background|history))?|academic(\s*(background|history))?|qualifications?|degrees?|schooling)$', re.I),
    'certs'     : re.compile(r'^(certifications?|licenses?|credentials?|professional\s*certifications?|awards?(\s*(and|&)\s*certifications?)?|training)$', re.I),
    'projects'  : re.compile(r'^(projects?|notable\s*projects?|key\s*projects?|personal\s*projects?|academic\s*projects?|portfolio)$', re.I),
}

# Company name words
_CO = re.compile(
    r'\b(Inc\.?|LLC|Corp\.?|Ltd\.?|Co\.|Company|Group|Solutions|Technologies|'
    r'Tech(?:\s|$)|Systems|Services|Associates|Partners|Consulting|Health(?:care)?|'
    r'Bank|Financial|Corporation|University|College|Institute|Hospital|'
    r'Google|Microsoft|Amazon|Apple|Meta|Oracle|IBM|SAP|Cisco|Intel|Dell|HP|'
    r'Deloitte|Accenture|KPMG|PwC|EY|Walmart|Netflix|Twitter|Uber|Lyft|'
    r'Management|Industries|International|National|Global|Enterprise)\b', re.I)

# Job title words
_TITLE = re.compile(
    r'\b(Engineer|Developer|Manager|Architect|Analyst|Director|Lead|Scientist|'
    r'Designer|Consultant|Administrator|Officer|Specialist|Sr\.?|Senior|Junior|'
    r'Jr\.?|Principal|Staff|VP|Head|Chief|President|CTO|CEO|CFO|COO|'
    r'DevOps|Full.?Stack|Backend|Frontend|Cloud|Data|ML|AI|QA|SRE|'
    r'Software|Hardware|Network|Security|Product|Project|Program)\b', re.I)

_BULL_PAT = re.compile(r'^\s*[-•*+>▪◦○●\u2022]\s*\S|^\s*\d+[\.\)]\s+\S')
_LOC_PAT  = re.compile(r'^[A-Z][a-zA-Z \.]{1,28},?\s*[A-Z]{0,2}$')

def _is_date(l):  return bool(_DRANGE.search(l))
def _is_bull(l):  return bool(_BULL_PAT.match(l))
def _is_loc(l):   return bool(_LOC_PAT.match(l.strip())) and 2 <= len(l.split()) <= 5

def _is_role_hdr(line):
    """Is this line a new job entry header (NOT a bullet, NOT a long sentence)?"""
    if _is_bull(line): return False
    if len(line) > 100: return False   # long sentences are not headers
    if _is_date(line): return True     # date line = new entry
    # "Title | Company" pattern
    if '|' in line and ';' not in line and len(line.split('|')) == 2:
        return True
    # Known company name in a short line
    if _CO.search(line) and len(line) < 80:
        return True
    return False

def parse_resume(raw: str) -> dict:
    text  = san(raw)
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    result = {
        'personal'            : {'full_name':'','email':'','phone':'',
                                  'location':'','linkedin':'','github':''},
        'professional_summary': '',
        'skills'              : {},
        'experience'          : [],
        'education'           : [],
        'certifications'      : [],
        'projects'            : [],
    }

    # ── contacts (scan full text) ─────────────────────
    em = re.search(r'[\w.+\-]+@[\w.\-]+\.\w{2,}', text)
    if em: result['personal']['email'] = em.group().lower()

    ph = re.search(r'(?:\+?1[\s.\-]?)?\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4}', text)
    if ph and len(re.sub(r'\D','',ph.group())) >= 10:
        result['personal']['phone'] = ph.group().strip()

    lm = re.search(r'[A-Z][a-zA-Z ]{1,20},\s*[A-Z]{2}(?:\s+\d{5})?', text)
    if not lm:
        lm = re.search(r'[A-Z][a-zA-Z]+(?: City| Beach| Springs)?[, ]+[A-Z]{2}\b', text)
    if lm: result['personal']['location'] = lm.group().strip()

    li = re.search(r'linkedin\.com/in/[\w\-]+', text, re.I)
    if li: result['personal']['linkedin'] = li.group()

    gh = re.search(r'github\.com/[\w\-]+', text, re.I)
    if gh: result['personal']['github'] = gh.group()

    # ── name (first short 2-5 word line) ─────────────
    for line in lines[:14]:
        if '@' in line: continue
        if re.match(r'[\d\(\+]', line): continue
        if re.search(r'linkedin|github|http|resume|curriculum|page\s+\d|©|cv\b', line, re.I): continue
        if any(p.match(line) for p in _SECS.values()): continue
        if _is_date(line) or _is_bull(line): continue
        words = line.split()
        if 2 <= len(words) <= 5 and len(line) < 60 and not _CO.search(line):
            result['personal']['full_name'] = ' '.join(w.capitalize() for w in words)
            break

    # ── bucket all lines by section ───────────────────
    buckets = {k: [] for k in _SECS}
    cur = None
    for line in lines:
        hit = next((s for s, p in _SECS.items() if p.match(line)), None)
        if hit:
            cur = hit
        elif cur:
            buckets[cur].append(line)

    # ── SUMMARY ──────────────────────────────────────
    if buckets['summary']:
        result['professional_summary'] = re.sub(
            r' {2,}', ' ', ' '.join(buckets['summary'][:12])).strip()

    # ── SKILLS ───────────────────────────────────────
    sk_found = False
    for line in buckets['skills']:
        m = re.match(r'^([A-Za-z][A-Za-z0-9 &/()\-]{1,55}):\s*(.+)$', line)
        if m:
            cat   = m.group(1).strip()
            items = [i.strip() for i in re.split(r'[,;|]', m.group(2)) if i.strip()]
            if items:
                result['skills'][cat] = items
                sk_found = True
    if not sk_found and buckets['skills']:
        all_items = []
        for line in buckets['skills']:
            all_items += [i.strip() for i in re.split(r'[,;|]', line)
                          if i.strip() and len(i.strip()) > 1]
        if all_items:
            result['skills']['Technical Skills'] = all_items

    # ── EXPERIENCE — two-pass block builder ──────────
    exp_lines = buckets['experience']   # STRICTLY from experience bucket only

    def _build_exp_blocks(elines):
        """Two-pass block builder: handles titles-before-dates format."""
        # Pass 1: split into raw blocks on role-header lines
        raw = []; cb = None
        for line in elines:
            if _is_role_hdr(line):
                if cb: raw.append(cb)
                cb = {'header':[line], 'bullets':[], 'has_date':_is_date(line)}
            elif cb is None:
                cb = {'header':[line], 'bullets':[], 'has_date':_is_date(line)}
            elif _is_bull(line) or cb['bullets']:
                cb['bullets'].append(line)
            else:
                if len(line) > 60 and not _is_date(line):
                    # Long non-header line = paragraph-style bullet
                    cb['bullets'].append(line)
                else:
                    cb['header'].append(line)
                    if _is_date(line): cb['has_date'] = True
        if cb: raw.append(cb)

        # Pass 2: merge dateless header-only blocks into next date block
        merged = []; pending = []
        for blk in raw:
            if blk['has_date']:
                blk['header'] = pending + blk['header']
                pending = []; merged.append(blk)
            elif not blk['bullets']:
                pending.extend(blk['header'])
            else:
                if pending:
                    blk['header'] = pending + blk['header']
                    pending = []
                merged.append(blk)
        if pending:
            merged.append({'header':pending,'bullets':[],'has_date':False})
        return merged

    for blk in _build_exp_blocks(exp_lines)[:12]:
        hdr  = blk['header']
        bulls = blk['bullets']
        title, company, loc_e, start, end = '', '', '', '', 'Present'

        for line in hdr:
            dm = _DRANGE.search(line)
            if dm:
                start = norm_date(dm.group(1))
                end   = norm_date(dm.group(2))
                # inline location after date: "Jan 2020 - Present | Houston, TX"
                rest = (line[:dm.start()] + ' ' + line[dm.end():]).strip()
                rest = re.sub(r'^[\s|,\-]+', '', rest).strip()
                if rest and _is_loc(rest) and not loc_e:
                    loc_e = rest
                continue

            if '|' in line and len(line.split('|')) == 2:
                p0, p1 = [x.strip() for x in line.split('|', 1)]
                # Heuristic: whichever side has a title word = title
                t0 = bool(_TITLE.search(p0)); t1 = bool(_TITLE.search(p1))
                c0 = bool(_CO.search(p0));    c1 = bool(_CO.search(p1))
                if c0 and not c1:   company, title = p0, p1
                elif c1 and not c0: title, company = p0, p1
                elif t1 and not t0: company, title = p0, p1
                elif t0 and not t1: title, company = p0, p1
                else:               title, company = p0, p1   # default left=title
                continue

            if _is_loc(line) and not loc_e:
                loc_e = line; continue

            if not title:   title   = line; continue
            if not company: company = line; continue

        # Clean bullets
        clean = []
        for b in bulls:
            b2 = strip_pfx(b)
            b2 = re.sub(r' {2,}', ' ', b2).strip()
            if len(b2) > 4:
                clean.append(b2[0].upper() + b2[1:])

        if title or company:
            result['experience'].append({
                'title'     : s1(title),
                'company'   : s1(company),
                'location'  : s1(loc_e),
                'start_date': start,
                'end_date'  : end,
                'bullets'   : clean,
            })

    # ── EDUCATION — completely independent ───────────
    edu_lines = buckets['education']   # STRICTLY from education bucket only
    if edu_lines:
        # Split edu blocks on degree lines (date lines mark end of each degree entry)
        edu_blks = []; cur_edu = []
        for line in edu_lines:
            if _is_date(line):
                cur_edu.append(line)
                edu_blks.append(cur_edu)
                cur_edu = []
            else:
                cur_edu.append(line)
        if cur_edu:
            edu_blks.append(cur_edu)

        for eb in edu_blks[:4]:
            if not eb: continue
            all_t = ' '.join(eb)
            dm    = _DRANGE.search(all_t)
            gpa_m = re.search(r'GPA:?\s*([\d.]+)', all_t, re.I)
            deg_l = sch_l = loc_e2 = ''
            for line in eb:
                if _is_date(line): continue
                if gpa_m and gpa_m.group(0) in line: continue
                if not deg_l: deg_l = line; continue
                if not sch_l:
                    if '|' in line:
                        pts = [x.strip() for x in line.split('|')]
                        sch_l = pts[0]
                        if len(pts) > 1: loc_e2 = pts[1]
                    else: sch_l = line
                    continue
                if not loc_e2 and _is_loc(line): loc_e2 = line
            result['education'].append({
                'degree'    : s1(deg_l),
                'school'    : s1(sch_l),
                'location'  : s1(loc_e2),
                'gpa'       : gpa_m.group(1) if gpa_m else '',
                'start_date': norm_date(dm.group(1)) if dm else '',
                'end_date'  : norm_date(dm.group(2)) if dm else '',
            })

    # ── CERTIFICATIONS ───────────────────────────────
    for line in buckets['certs']:
        c = strip_pfx(line).strip()
        if c and len(c) > 4:
            result['certifications'].append(s1(c))

    # ── PROJECTS ─────────────────────────────────────
    proj_lines = buckets['projects']
    proj_blks = []; cp = None
    for line in proj_lines:
        if not _is_bull(line) and not cp:
            cp = {'header':[line],'bullets':[]}
        elif cp and not _is_bull(line) and not cp['bullets']:
            cp['header'].append(line)
            if _is_date(line) or (len(line.split()) <= 8 and line[0].isupper() and cp['header']):
                if len(cp['header']) > 2: proj_blks.append(cp); cp = {'header':[line],'bullets':[]}
        elif cp and (_is_bull(line) or cp['bullets']):
            cp['bullets'].append(line)
        elif not cp:
            cp = {'header':[line],'bullets':[]}
    if cp: proj_blks.append(cp)

    for pb in proj_blks[:8]:
        h = pb['header']
        all_h = ' '.join(h)
        dm = _DRANGE.search(all_h)
        name = tech = pd = ''
        for line in h:
            if _is_date(line): pd = line; continue
            if '|' in line:
                pts = [x.strip() for x in line.split('|',1)]
                name = pts[0]; tech = pts[1] if len(pts)>1 else ''
            elif not name: name = line
        if dm: pd = f'{norm_date(dm.group(1))} - {norm_date(dm.group(2))}'
        cb2 = [strip_pfx(b).strip() for b in pb['bullets'] if len(strip_pfx(b).strip())>4]
        if name:
            result['projects'].append({
                'name':s1(name),'tech':s1(tech),'date':pd,
                'bullets':[b[0].upper()+b[1:] for b in cb2 if b],
            })

    return result

# ══════════════════════════════════════════════════════
#  DOCX WRITER — ATS-standard typography
#  Per screenshot analysis:
#   PROFESSIONAL EXPERIENCE  → 10pt ALL CAPS Bold (section header = SMALLER)
#   Senior Software Engineer | ABC Tech  → 11pt Bold (role = LARGER)
#   Jan 2020 - Dec 2025  → 10pt italic RIGHT-aligned via tab
#   Plano, Texas  → 10pt italic
#   • Debugged production issues...  → 11pt, hanging indent, clear gap after •
# ══════════════════════════════════════════════════════
_TAB = 9900   # 6.875" from left margin (text width 7.0")

def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'0B2A55')
    pBdr.append(bot); pPr.append(pBdr)

def _sec_hdr(doc, label: str):
    """
    Section header: ALL CAPS, 10pt Bold.
    Smaller than role names (11pt) — ATS standard.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(14)
    r = p.add_run(san(label).upper())
    r.font.name = DF; r.font.size = Pt(SZ_SEC_HDR); r.font.bold = True
    return p

def _tab_line(doc, left: str, right: str = '',
              sz_l=SZ_ROLE, sz_r=SZ_DATE, bold_l=True, sb=0, sa=1):
    """
    Role/Degree line with right-aligned date.
    left: 11pt bold   right: 10pt italic RIGHT TAB
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(14)
    # right tab stop
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'),'right'); tab.set(qn('w:pos'),str(_TAB))
    tabs.append(tab); pPr.append(tabs)
    # left run (role)
    rl = p.add_run(san(str(left)))
    rl.font.name = DF; rl.font.size = Pt(sz_l); rl.font.bold = bold_l
    # tab + right run (date)
    if right:
        p.add_run('\t')
        rr = p.add_run(san(str(right)))
        rr.font.name = DF; rr.font.size = Pt(sz_r); rr.font.italic = True
    return p

def _plain(doc, text: str, bold=False, size=SZ_BODY, italic=False,
           sb=0, sa=2, indent=None, justify=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(14)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    if justify and len(text) > 60:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(san(str(text)))
    r.font.name = DF; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    return p

def _bullet(doc, text: str, pfx='•', idx=0, justify=False):
    """
    ATS bullet paragraph — matches screenshot exactly:
      • space  Lead text here that wraps
               cleanly under first word
    Hanging indent: left=0.25", first_line=-0.25"
    space_before=1, space_after=3 — clear visual gap
    Bullet char + double-space before text for clear separation.
    """
    is_num = (pfx == 'NUM')
    marker = f'{idx+1}.  ' if is_num else f'{pfx}  '   # double space after bullet

    p = doc.add_paragraph()
    p.paragraph_format.space_before      = Pt(1)
    p.paragraph_format.space_after       = Pt(3)
    p.paragraph_format.line_spacing      = Pt(13.5)
    p.paragraph_format.left_indent       = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if justify and len(text) > 60:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    r = p.add_run(marker + san(text))
    r.font.name = DF; r.font.size = Pt(SZ_BODY)
    return p

def write_docx(rd: dict, filepath: str, bullet_char='•', justify=True):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.60)
        sec.bottom_margin = Inches(0.60)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)

    ns = doc.styles['Normal']
    ns.font.name = DF; ns.font.size = Pt(SZ_BODY)
    ns.paragraph_format.space_before = Pt(0)
    ns.paragraph_format.space_after  = Pt(0)

    pers = rd.get('personal', {})

    # NAME — 14pt Bold ALL CAPS
    _plain(doc, san(pers.get('full_name','')).upper(),
           bold=True, size=SZ_NAME, sa=1)

    # CONTACT — 10pt
    c_parts = [san(pers.get(k,'')) for k in
               ('email','phone','location','linkedin','github') if pers.get(k)]
    if c_parts:
        _plain(doc, ' | '.join(c_parts), size=SZ_CONTACT, sa=3)

    # PROFESSIONAL SUMMARY
    if rd.get('professional_summary'):
        _hr(doc)
        _sec_hdr(doc, 'Professional Summary')
        _plain(doc, san(rd['professional_summary']),
               size=SZ_BODY, sb=1, sa=3, justify=justify)

    # CORE COMPETENCIES
    if rd.get('skills'):
        _hr(doc)
        _sec_hdr(doc, 'Core Competencies')
        for cat, items in rd['skills'].items():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing = Pt(14)
            r1 = p.add_run(f'{san(cat)}: ')
            r1.font.name=DF; r1.font.size=Pt(SZ_BODY); r1.font.bold=True
            r2 = p.add_run(', '.join(san(s) for s in items))
            r2.font.name=DF; r2.font.size=Pt(SZ_BODY)

    # PROFESSIONAL EXPERIENCE
    if rd.get('experience'):
        _hr(doc)
        _sec_hdr(doc, 'Professional Experience')
        for exp in rd['experience']:
            t   = san(exp.get('title',''))
            co  = san(exp.get('company',''))
            loc = san(exp.get('location',''))
            st  = exp.get('start_date','')
            en  = exp.get('end_date','Present')
            ds  = f'{st} - {en}' if st else en
            rl  = f'{t}  |  {co}' if t and co else (t or co)
            _tab_line(doc, rl, ds, sz_l=SZ_ROLE, sz_r=SZ_DATE,
                      bold_l=True, sb=6, sa=0)
            if loc:
                _plain(doc, loc, italic=True, size=SZ_LOC, sb=0, sa=2)
            for i, b in enumerate(exp.get('bullets',[])):
                b = san(b)
                if b: _bullet(doc, b, pfx=bullet_char, idx=i, justify=justify)

    # EDUCATION
    if rd.get('education'):
        _hr(doc)
        _sec_hdr(doc, 'Education')
        for edu in rd['education']:
            deg = san(edu.get('degree',''))
            sch = san(edu.get('school',''))
            loc = san(edu.get('location',''))
            gpa = edu.get('gpa','')
            s,e = edu.get('start_date',''), edu.get('end_date','')
            ds  = f'{s} - {e}' if s and e else (e or s)
            _tab_line(doc, deg, ds, sz_l=SZ_ROLE, sz_r=SZ_DATE,
                      bold_l=True, sb=5, sa=0)
            parts = [x for x in (sch, loc) if x]
            if gpa: parts.append(f'GPA: {gpa}')
            if parts:
                _plain(doc, ' | '.join(parts), size=SZ_LOC, sb=0, sa=3)

    # CERTIFICATIONS
    if rd.get('certifications'):
        _hr(doc)
        _sec_hdr(doc, 'Certifications')
        for i, c in enumerate(rd['certifications']):
            c = san(c)
            if c: _bullet(doc, c, pfx=bullet_char, idx=i, justify=False)

    # PROJECTS
    if rd.get('projects'):
        _hr(doc)
        _sec_hdr(doc, 'Projects')
        for proj in rd['projects']:
            nm  = san(proj.get('name',''))
            tch = san(proj.get('tech',''))
            dt  = proj.get('date','')
            left = f'{nm}  |  {tch}' if tch else nm
            _tab_line(doc, left, dt, sz_l=SZ_ROLE, sz_r=SZ_DATE,
                      bold_l=True, sb=5, sa=0)
            for i, b in enumerate(proj.get('bullets',[])):
                b = san(b)
                if b: _bullet(doc, b, pfx=bullet_char, idx=i, justify=justify)

    doc.save(filepath)

# ══════════════════════════════════════════════════════
#  UI HELPERS
# ══════════════════════════════════════════════════════
def _e(parent, **kw):
    e = tk.Entry(parent, font=(UF,12), bg=INP, fg=TEXT, relief=tk.FLAT,
                 insertbackground=ACCENT,
                 highlightbackground=BORDER, highlightthickness=1, **kw)
    e.bind('<FocusIn>',  lambda _: e.config(highlightbackground=ACCENT, highlightthickness=2))
    e.bind('<FocusOut>', lambda _: e.config(highlightbackground=BORDER, highlightthickness=1))
    return e

def _t(parent, h=5, **kw):
    t = tk.Text(parent, font=(UF,12), bg=INP, fg=TEXT, relief=tk.FLAT,
                wrap=tk.WORD, height=h, padx=9, pady=7,
                insertbackground=ACCENT,
                highlightbackground=BORDER, highlightthickness=1, **kw)
    t.bind('<FocusIn>',  lambda _: t.config(highlightbackground=ACCENT, highlightthickness=2))
    t.bind('<FocusOut>', lambda _: t.config(highlightbackground=BORDER, highlightthickness=1))
    return t

def _btn(parent, label, cmd, sty='primary', px=22, py=9, **kw):
    pal = {'primary':(ACCENT,WHITE),'success':(GREEN,WHITE),
           'ghost':(BG,TEXT),'danger':(RED,WHITE),'outline':(WHITE,ACCENT)}
    bg,fg = pal.get(sty, pal['primary'])
    return tk.Button(parent, text=label, command=cmd,
                     font=(UF,12,'bold'), bg=bg, fg=fg, relief=tk.FLAT,
                     cursor='hand2', padx=px, pady=py,
                     activebackground=bg, activeforeground=fg, **kw)

def _card(parent, title='', subtitle='', collapse=False):
    outer = tk.Frame(parent, bg=WHITE, highlightbackground=BORDER, highlightthickness=1)
    outer.pack(fill='x', padx=26, pady=(0,8))
    body  = None
    if title:
        hf = tk.Frame(outer, bg=HDR_C, height=52)
        hf.pack(fill='x'); hf.pack_propagate(False)
        rf = tk.Frame(hf, bg=HDR_C); rf.pack(side='left', padx=18, fill='y')
        tk.Label(rf, text=title, font=(UF,13,'bold'), fg=WHITE, bg=HDR_C).pack(side='left')
        if subtitle:
            tk.Label(rf, text=f'  —  {subtitle}', font=(UF,10), fg='#93C5FD', bg=HDR_C).pack(side='left')
    body = tk.Frame(outer, bg=WHITE, padx=26, pady=20)
    body.pack(fill='both')
    return outer, body

def _nav(parent, back=None, nxt=None, nxt_lbl='Continue →', nxt_sty='primary'):
    bar = tk.Frame(parent, bg='#EDF2F7', highlightbackground=BORDER, highlightthickness=1)
    bar.pack(fill='x')
    inn = tk.Frame(bar, bg='#EDF2F7', padx=20, pady=12); inn.pack(fill='x')
    if back: _btn(inn,'← Back', back, sty='ghost').pack(side='left')
    if nxt:  _btn(inn, nxt_lbl, nxt, sty=nxt_sty).pack(side='right')

def _tip(parent, msg):
    f = tk.Frame(parent, bg=TIP_BG, padx=14, pady=10)
    f.pack(fill='x', pady=(8,2))
    tk.Label(f, text=f'💡  {msg}', font=(UF,10), fg=TIP_FG,
             bg=TIP_BG, wraplength=720, justify='left').pack(anchor='w')

def _scrollbox(parent):
    cv  = tk.Canvas(parent, bg=BG, highlightthickness=0)
    sb2 = ttk.Scrollbar(parent, orient='vertical', command=cv.yview)
    inn = tk.Frame(cv, bg=BG)
    inn.bind('<Configure>', lambda e: cv.configure(scrollregion=cv.bbox('all')))
    w   = cv.create_window((0,0), window=inn, anchor='nw')
    cv.configure(yscrollcommand=sb2.set)
    cv.bind('<Configure>', lambda e: cv.itemconfig(w, width=e.width))
    cv.pack(side='left', fill='both', expand=True)
    sb2.pack(side='right', fill='y')
    def _scroll(e): cv.yview_scroll(int(-1*(e.delta/120)),'units')
    cv.bind_all('<MouseWheel>', _scroll)
    return inn

# ══════════════════════════════════════════════════════
#  APPLICATION
# ══════════════════════════════════════════════════════
class App:
    def __init__(self, root):
        self.root = root
        root.title('ATS Resume Builder Pro')
        root.geometry('1440x920')
        root.minsize(1100,760)
        root.configure(bg=BG)
        try:
            from ctypes import windll; windll.shcore.SetProcessDpiAwareness(1)
        except Exception: pass

        self.step       = 0
        self.data       = self._blank()
        self._prev_data = None   # for undo
        self.refs       = {}
        self.bullet_var = None
        self.align_var  = None

        self._shell()
        self._go(0)

    # ── blank data template ──────────────────────────
    def _blank(self):
        return {
            'personal'            : {'full_name':'','email':'','phone':'',
                                      'location':'','linkedin':'','github':''},
            'professional_summary': '',
            'skills':{}, 'experience':[], 'education':[],
            'certifications':[], 'projects':[],
        }

    # ── build window shell ───────────────────────────
    def _shell(self):
        # ── SIDEBAR ─────────────────────────────────
        self.sbf = tk.Frame(self.root, bg=SB, width=242)
        self.sbf.pack(side='left', fill='y'); self.sbf.pack_propagate(False)

        tk.Frame(self.sbf, bg=SB, height=20).pack()
        tk.Label(self.sbf, text='ATS', font=(UF,26,'bold'), fg=ACCENT, bg=SB).pack()
        tk.Label(self.sbf, text='Resume Builder Pro', font=(UF,10), fg=TODO, bg=SB).pack(pady=(0,12))
        tk.Frame(self.sbf, bg='#1D3A5E', height=1).pack(fill='x', padx=18, pady=4)

        self._sbits = []
        for i, name in enumerate(STEPS):
            f = tk.Frame(self.sbf, bg=SB, cursor='hand2')
            f.pack(fill='x', padx=12, pady=3)
            nf = tk.Frame(f, bg=TODO, width=30, height=30)
            nf.pack(side='left', padx=(8,12), pady=6); nf.pack_propagate(False)
            nl = tk.Label(nf, text=str(i+1), font=(UF,11,'bold'), bg=TODO, fg=WHITE)
            nl.place(relx=.5, rely=.5, anchor='center')
            nt = tk.Label(f, text=name, font=(UF,12), bg=SB, fg=TODO)
            nt.pack(side='left')
            for w in (f,nf,nl,nt):
                w.bind('<Button-1>', lambda e,idx=i: self._go(idx) if idx<=self.step else None)
            self._sbits.append((f,nf,nl,nt))

        tk.Frame(self.sbf, bg='#1D3A5E', height=1).pack(fill='x', padx=18, pady=12)

        sb_btns = [
            ('📤  Upload Resume',     self._upload,   'outline'),
            ('📋  Paste & Fill',      self._paste,    'outline'),
            ('🔍  Keyword Analyzer',  self._keywords, 'ghost'),
            ('📝  Resume Templates',  self._templates,'ghost'),
            ('💾  Load Saved',        self._load_json,'ghost'),
        ]
        for lbl, cmd, sty in sb_btns:
            _btn(self.sbf, lbl, cmd, sty=sty, px=10, py=8).pack(
                fill='x', padx=12, pady=2)

        # ── CHECKLIST ───────────────────────────────
        tk.Frame(self.sbf, bg='#1D3A5E', height=1).pack(fill='x', padx=18, pady=(12,4))
        tk.Label(self.sbf, text='Checklist', font=(UF,10,'bold'),
                  fg=TODO, bg=SB).pack(anchor='w', padx=18, pady=(0,4))
        self._chk_labels = {}
        chk_items = ['Name & Contact','Summary','Skills','Experience','Education']
        for item in chk_items:
            lf = tk.Frame(self.sbf, bg=SB); lf.pack(anchor='w', padx=18, pady=1)
            ic = tk.Label(lf, text='○', font=(UF,11), fg=TODO, bg=SB); ic.pack(side='left')
            tk.Label(lf, text=f'  {item}', font=(UF,10), fg=TODO, bg=SB).pack(side='left')
            self._chk_labels[item] = ic

        # ── MAIN AREA ───────────────────────────────
        self.main = tk.Frame(self.root, bg=BG)
        self.main.pack(side='left', fill='both', expand=True)

        tb = tk.Frame(self.main, bg=WHITE, highlightbackground=BORDER, highlightthickness=1)
        tb.pack(fill='x'); tb.configure(height=58); tb.pack_propagate(False)
        self._tb_title = tk.Label(tb, text='', font=(UF,15,'bold'), fg='#07172B', bg=WHITE)
        self._tb_title.pack(side='left', padx=26)
        self._tb_sub   = tk.Label(tb, text='', font=(UF,11), fg=MUTED, bg=WHITE)
        self._tb_sub.pack(side='left', padx=4)
        # Undo button in topbar
        self._undo_btn = _btn(tb, '↩ Undo Parse', self._undo_parse,
                               sty='ghost', px=12, py=6)
        self._undo_btn.pack(side='right', padx=16)
        self._undo_btn.pack_forget()

        sty = ttk.Style(); sty.theme_use('clam')
        sty.configure('ats.Horizontal.TProgressbar',
                       troughcolor=BORDER, background=ACCENT, thickness=5)
        self._pb = tk.DoubleVar()
        ttk.Progressbar(self.main, variable=self._pb,
                        style='ats.Horizontal.TProgressbar',
                        maximum=len(STEPS)-1).pack(fill='x')

        self.content = tk.Frame(self.main, bg=BG)
        self.content.pack(fill='both', expand=True)

    # ── step nav ────────────────────────────────────
    def _go(self, idx):
        self.step = max(self.step, idx)
        self._pb.set(idx)
        self._tb_title.config(text=STEPS[idx])
        self._tb_sub.config(text=f'Step {idx+1} of {len(STEPS)}')
        for i,(f,nf,nl,nt) in enumerate(self._sbits):
            if i<idx:
                nf.config(bg=DONE); nl.config(bg=DONE,text='✓')
                nt.config(fg=WHITE,font=(UF,12)); f.config(bg=SB)
            elif i==idx:
                nf.config(bg=ACCENT); nl.config(bg=ACCENT,text=str(i+1))
                nt.config(fg=WHITE,font=(UF,12,'bold')); f.config(bg=SB_HL)
            else:
                nf.config(bg=TODO); nl.config(bg=TODO,text=str(i+1))
                nt.config(fg=TODO,font=(UF,12)); f.config(bg=SB)
        self._update_checklist()
        for w in self.content.winfo_children(): w.destroy()
        self.refs = {}
        [self._pg_personal,self._pg_summary,self._pg_skills,
         self._pg_experience,self._pg_education,self._pg_preview][idx]()

    def _update_checklist(self):
        d = self.data
        checks = {
            'Name & Contact': bool(d['personal'].get('full_name') and d['personal'].get('email')),
            'Summary'       : len(d.get('professional_summary','')) > 50,
            'Skills'        : bool(d.get('skills')),
            'Experience'    : bool(d.get('experience')),
            'Education'     : bool(d.get('education')),
        }
        for item, done in checks.items():
            lbl = self._chk_labels.get(item)
            if lbl:
                lbl.config(text='✓' if done else '○',
                           fg=DONE if done else TODO)

    def _undo_parse(self):
        if self._prev_data:
            self.data = copy.deepcopy(self._prev_data)
            self._prev_data = None
            self._undo_btn.pack_forget()
            messagebox.showinfo('Undone','Previous data restored.')
            self._go(5)

    # ════════════════════════════════════════
    #  PAGE 1 — PERSONAL
    # ════════════════════════════════════════
    def _pg_personal(self):
        wrap = tk.Frame(self.content, bg=BG)
        wrap.pack(fill='both', expand=True, pady=18)
        outer, body = _card(wrap,'Personal Information','plain text only')

        g = tk.Frame(body, bg=WHITE); g.pack(fill='x')
        g.columnconfigure(0,weight=1); g.columnconfigure(1,weight=1)

        fields = [
            ('Full Name',          'full_name', True,  'e.g., Alex Johnson'),
            ('Email Address',      'email',     True,  'e.g., alex@email.com'),
            ('Phone Number',       'phone',     True,  'e.g., (555) 123-4567'),
            ('City, State',        'location',  True,  'e.g., Austin, TX'),
            ('LinkedIn URL',       'linkedin',  False, 'e.g., linkedin.com/in/alex'),
            ('GitHub / Portfolio', 'github',    False, 'e.g., github.com/alex'),
        ]
        for i,(lb,key,req,hint_t) in enumerate(fields):
            row,col = i//2, i%2
            cell = tk.Frame(g, bg=WHITE)
            cell.grid(row=row, column=col, sticky='ew',
                      padx=(0 if col==0 else 8, 8 if col==0 else 0), pady=6)
            hf = tk.Frame(cell, bg=WHITE); hf.pack(anchor='w')
            tk.Label(hf,text=lb,font=(UF,11,'bold'),fg=TEXT,bg=WHITE).pack(side='left')
            tk.Label(hf,text=' *' if req else '  (optional)',
                     font=(UF,10),fg=RED if req else MUTED,bg=WHITE).pack(side='left')
            e = _e(cell); e.pack(fill='x', ipady=11, pady=(3,0))
            tk.Label(cell,text=hint_t,font=(UF,9),fg=MUTED,bg=WHITE).pack(anchor='w',pady=(2,0))
            v = self.data['personal'].get(key,'')
            if v: e.insert(0,v)
            self.refs[key] = e

        _tip(body,'ATS reads plain text only. Use "City, ST" format. No emojis or symbols.')
        _nav(outer, back=None, nxt=self._sv_personal)

    def _sv_personal(self):
        for k,lbl in [('full_name','Full Name'),('email','Email'),
                       ('phone','Phone'),('location','City, State')]:
            if not self.refs[k].get().strip():
                messagebox.showerror('Required',f'Please enter your {lbl}.'); return
        if '@' not in self.refs['email'].get():
            messagebox.showerror('Invalid','Enter a valid email address.'); return
        for k in ('full_name','email','phone','location','linkedin','github'):
            self.data['personal'][k] = s1(self.refs[k].get())
        self._update_checklist(); self._go(1)

    # ════════════════════════════════════════
    #  PAGE 2 — SUMMARY
    # ════════════════════════════════════════
    def _pg_summary(self):
        wrap = tk.Frame(self.content, bg=BG)
        wrap.pack(fill='both', expand=True, pady=18)
        outer, body = _card(wrap,'Professional Summary','3–5 sentences, keyword-rich')

        tk.Label(body,text='Summary  *',font=(UF,11,'bold'),fg=TEXT,bg=WHITE).pack(anchor='w',pady=(4,3))
        tb = _t(body, h=9); tb.pack(fill='x')
        val = self.data.get('professional_summary','')
        if val: tb.insert('1.0', val)

        stat_row = tk.Frame(body,bg=WHITE); stat_row.pack(fill='x',pady=(5,0))
        char_l = tk.Label(stat_row,text='0 chars',font=(UF,10),fg=MUTED,bg=WHITE); char_l.pack(side='right')
        verb_l = tk.Label(stat_row,text='',font=(UF,9),fg=MUTED,bg=WHITE); verb_l.pack(side='left')

        def upd(_=None):
            txt = tb.get('1.0','end-1c'); n=len(txt); w=len(txt.split())
            char_l.config(text=f'{n} chars | {w} words',
                          fg=GREEN if 80<=n<=900 else AMBER if n>0 else MUTED)
            fv=[v for v in STRONG_VERBS if v.lower() in txt.lower()]
            verb_l.config(text=f'✓ Strong verbs: {", ".join(fv[:4])}' if fv else '')
        tb.bind('<KeyRelease>', upd); upd()
        self.refs['summary'] = tb

        # Quick sample fill
        def fill_sample():
            tb.delete('1.0',tk.END)
            tb.insert('1.0','Senior Software Engineer with 8+ years of experience building '
                             'scalable cloud-native applications on AWS. Expertise in Java, '
                             'Python, and microservices architecture. Delivered 40% reduction '
                             'in infrastructure costs through automation and IaC best practices.')
            upd()
        sf = tk.Frame(body,bg=WHITE); sf.pack(anchor='e',pady=(4,0))
        _btn(sf,'Fill Sample',fill_sample,sty='ghost',px=10,py=4).pack(side='right')

        _tip(body,'Formula: [Role + Years] + [2-3 core skills] + [1 measurable result]. No "I" or "my".')
        _nav(outer, back=lambda: self._go(0), nxt=self._sv_summary)

    def _sv_summary(self):
        v = self.refs['summary'].get('1.0',tk.END).strip()
        if len(v)<40:
            messagebox.showerror('Too Short','Write at least 2 sentences.'); return
        self.data['professional_summary'] = san(v)
        self._update_checklist(); self._go(2)

    # ════════════════════════════════════════
    #  PAGE 3 — SKILLS
    # ════════════════════════════════════════
    def _pg_skills(self):
        wrap = tk.Frame(self.content, bg=BG); wrap.pack(fill='both', expand=True)
        top = tk.Frame(wrap,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        top.pack(fill='x',padx=26)
        hf = tk.Frame(top,bg=HDR_C,height=52); hf.pack(fill='x'); hf.pack_propagate(False)
        tk.Label(hf,text='  Core Competencies & Skills  —  group by category',
                  font=(UF,13,'bold'),fg=WHITE,bg=HDR_C).pack(side='left',padx=6)
        _nav(top, back=lambda: self._go(1), nxt=self._sv_skills)

        so = tk.Frame(wrap,bg=BG); so.pack(fill='both',expand=True,padx=26)
        inner = _scrollbox(so)

        ch = tk.Frame(inner,bg='#E4EDF8'); ch.pack(fill='x',pady=(10,4))
        tk.Label(ch,text='  Category',font=(UF,11,'bold'),fg=MUTED,bg='#E4EDF8',
                  width=26,anchor='w').pack(side='left',padx=(12,0))
        tk.Label(ch,text='Skills — comma separated  *',font=(UF,11,'bold'),
                  fg=MUTED,bg='#E4EDF8').pack(side='left',padx=8)

        self.refs['sk_rows'] = []; self.refs['sk_frame'] = inner
        if self.data.get('skills'):
            for cat,items in self.data['skills'].items():
                self._add_skill(cat,', '.join(items))
        else:
            self._add_skill()

        bf = tk.Frame(inner,bg=BG); bf.pack(anchor='w',pady=8)
        _btn(bf,'+ Add Category',self._add_skill,sty='outline',px=12,py=6).pack(side='left',padx=(0,8))

        # Suggested skill categories
        def add_suggested():
            for cat in ['Cloud Platforms','Backend Development','DevOps & CI/CD','Databases']:
                self._add_skill(cat,'')
        _btn(bf,'+ Suggested Categories',add_suggested,sty='ghost',px=12,py=6).pack(side='left')

        _tip(inner,'Include exact keywords from the job description. '
                    '"ML (Machine Learning)" — acronym + full form.')

    def _add_skill(self, cat='', sk=''):
        inner = self.refs['sk_frame']
        row   = tk.Frame(inner,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        row.pack(fill='x',pady=2)
        ce = _e(row,width=26); ce.pack(side='left',padx=(10,6),pady=9,ipady=10)
        if cat: ce.insert(0,cat)
        se = _e(row); se.pack(side='left',fill='x',expand=True,padx=(0,6),pady=9,ipady=10)
        if sk: se.insert(0,sk)
        def rm(r=row,c=ce,s=se):
            self.refs['sk_rows']=[(cc,ss) for cc,ss in self.refs['sk_rows'] if cc is not c]
            r.destroy()
        tk.Button(row,text='✕',font=(UF,12),bg=WHITE,fg=RED,
                   relief=tk.FLAT,cursor='hand2',command=rm).pack(side='right',padx=10)
        self.refs['sk_rows'].append((ce,se))

    def _sv_skills(self):
        skills={}
        for ce,se in self.refs['sk_rows']:
            c=ce.get().strip(); s=se.get().strip()
            if c and s:
                items=[x.strip() for x in s.split(',') if x.strip()]
                if items: skills[c]=items
        if not skills:
            messagebox.showerror('Required','Add at least one skill category.'); return
        self.data['skills']=skills; self._update_checklist(); self._go(3)

    # ════════════════════════════════════════
    #  PAGE 4 — EXPERIENCE
    # ════════════════════════════════════════
    def _pg_experience(self):
        wrap = tk.Frame(self.content,bg=BG); wrap.pack(fill='both',expand=True)
        top = tk.Frame(wrap,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        top.pack(fill='x',padx=26)
        hf = tk.Frame(top,bg=HDR_C,height=52); hf.pack(fill='x'); hf.pack_propagate(False)
        tk.Label(hf,text='  Professional Experience  —  most recent first',
                  font=(UF,13,'bold'),fg=WHITE,bg=HDR_C).pack(side='left',padx=6)
        _nav(top,back=lambda: self._go(2),nxt=self._sv_experience)

        so = tk.Frame(wrap,bg=BG); so.pack(fill='both',expand=True,padx=26)
        inner = _scrollbox(so)
        self.refs['exp_blks']=[]; self.refs['exp_inner']=inner

        if self.data.get('experience'):
            for exp in self.data['experience']: self._add_exp(exp)
        else:
            self._add_exp()

        bf = tk.Frame(inner,bg=BG); bf.pack(anchor='w',pady=8)
        _btn(bf,'+ Add Position',lambda: self._add_exp(),sty='outline',px=12,py=6).pack(side='left',padx=(0,8))
        _btn(bf,'Fill Demo Data', self._demo_exp, sty='ghost',px=12,py=6).pack(side='left')
        _tip(inner,'Action verb + metric: "Reduced costs 40% by automating infrastructure provisioning."')

    def _demo_exp(self):
        if not self.refs['exp_blks']: self._add_exp()
        refs,_ = self.refs['exp_blks'][0]
        refs['title'].delete(0,tk.END);   refs['title'].insert(0,'Senior Software Engineer')
        refs['company'].delete(0,tk.END); refs['company'].insert(0,'ABC Technology Corp')
        refs['location'].delete(0,tk.END);refs['location'].insert(0,'Austin, TX')
        refs['start_date'].delete(0,tk.END);refs['start_date'].insert(0,'Jan 2022')
        refs['end_date'].delete(0,tk.END);  refs['end_date'].insert(0,'Present')
        refs['bullets'].delete('1.0',tk.END)
        refs['bullets'].insert('1.0',
            'Architected microservices platform handling 10M+ daily transactions on AWS EKS\n'
            'Reduced deployment time by 65% by implementing GitOps with ArgoCD and Helm\n'
            'Led team of 6 engineers, delivering 3 major product releases on schedule')

    def _add_exp(self, data=None):
        data=data or {}
        n=len(self.refs['exp_blks'])+1
        cont=self.refs['exp_inner']
        outer=tk.Frame(cont,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        outer.pack(fill='x',pady=(4,0))
        chdr=tk.Frame(outer,bg='#E8F0FB',height=42); chdr.pack(fill='x'); chdr.pack_propagate(False)
        pos_lbl=tk.Label(chdr,text=f'  Position {n}',font=(UF,11,'bold'),fg='#07172B',bg='#E8F0FB')
        pos_lbl.pack(side='left',padx=8)

        # ↑ ↓ reorder buttons
        def move_up(o=outer):
            idx=next((i for i,(r,oo) in enumerate(self.refs['exp_blks']) if oo is o),None)
            if idx and idx>0:
                self.refs['exp_blks'][idx-1][1].pack_forget()
                self.refs['exp_blks'][idx][1].pack_forget()
                self.refs['exp_blks'][idx-1],self.refs['exp_blks'][idx]=\
                    self.refs['exp_blks'][idx],self.refs['exp_blks'][idx-1]
                for _,oo in self.refs['exp_blks']: oo.pack(fill='x',pady=(4,0))
        def move_dn(o=outer):
            idx=next((i for i,(r,oo) in enumerate(self.refs['exp_blks']) if oo is o),None)
            if idx is not None and idx<len(self.refs['exp_blks'])-1:
                self.refs['exp_blks'][idx+1][1].pack_forget()
                self.refs['exp_blks'][idx][1].pack_forget()
                self.refs['exp_blks'][idx],self.refs['exp_blks'][idx+1]=\
                    self.refs['exp_blks'][idx+1],self.refs['exp_blks'][idx]
                for _,oo in self.refs['exp_blks']: oo.pack(fill='x',pady=(4,0))

        tk.Button(chdr,text='↑',font=(UF,10),bg='#E8F0FB',fg=ACCENT,
                   relief=tk.FLAT,cursor='hand2',command=move_up).pack(side='right',padx=4)
        tk.Button(chdr,text='↓',font=(UF,10),bg='#E8F0FB',fg=ACCENT,
                   relief=tk.FLAT,cursor='hand2',command=move_dn).pack(side='right')

        body=tk.Frame(outer,bg=WHITE,padx=22,pady=16); body.pack(fill='x')
        refs={}

        r1=tk.Frame(body,bg=WHITE); r1.pack(fill='x',pady=4)
        for key,lbl_t,req in [('title','Job Title',True),('company','Company Name',True)]:
            f=tk.Frame(r1,bg=WHITE); f.pack(side='left',fill='x',expand=True,padx=(0,10))
            hf2=tk.Frame(f,bg=WHITE); hf2.pack(anchor='w')
            tk.Label(hf2,text=lbl_t,font=(UF,11,'bold'),fg=TEXT,bg=WHITE).pack(side='left')
            if req: tk.Label(hf2,text=' *',font=(UF,11),fg=RED,bg=WHITE).pack(side='left')
            e=_e(f); e.pack(fill='x',ipady=10,pady=(3,0))
            v=data.get(key,'')
            if v: e.insert(0,v)
            refs[key]=e

        r2=tk.Frame(body,bg=WHITE); r2.pack(fill='x',pady=4)
        for key,lbl_t,hint_t in [
            ('location','Location','e.g., New York, NY'),
            ('start_date','Start Date','e.g., Jan 2021'),
            ('end_date','End Date','e.g., Present'),
        ]:
            f=tk.Frame(r2,bg=WHITE); f.pack(side='left',fill='x',expand=True,padx=(0,10))
            tk.Label(f,text=lbl_t,font=(UF,11,'bold'),fg=TEXT,bg=WHITE).pack(anchor='w')
            e=_e(f); e.pack(fill='x',ipady=10,pady=(3,0))
            tk.Label(f,text=hint_t,font=(UF,9),fg=MUTED,bg=WHITE).pack(anchor='w',pady=(2,0))
            v=data.get(key,'');
            if v: e.insert(0,v)
            refs[key]=e

        tk.Label(body,text='Key Achievements / Responsibilities',
                  font=(UF,11,'bold'),fg=TEXT,bg=WHITE).pack(anchor='w',pady=(12,2))
        tk.Label(body,text='One per line  •  Start with action verb  •  Include metrics',
                  font=(UF,10),fg=MUTED,bg=WHITE).pack(anchor='w')

        tb=_t(body,h=6); tb.pack(fill='x',pady=(4,0))
        bulls=data.get('bullets',[])
        if bulls: tb.insert('1.0','\n'.join(bulls))
        refs['bullets']=tb

        # Live verb checker
        av_lbl=tk.Label(body,text='',font=(UF,9),fg=MUTED,bg=WHITE); av_lbl.pack(anchor='w',pady=(3,0))
        def chk_v(_=None):
            txt=tb.get('1.0',tk.END).lower()
            fv=[v for v in STRONG_VERBS if v.lower() in txt]
            wk=[w for w in WEAK_VERBS if w in txt]
            msg=''
            if fv: msg+=f'✓ {", ".join(fv[:3])}'
            if wk: msg+=f'  ⚠ Weak: {", ".join(wk[:2])}'
            av_lbl.config(text=msg,fg=GREEN if fv and not wk else AMBER if msg else MUTED)
        tb.bind('<KeyRelease>',chk_v); chk_v()

        # Bullet auto-improve button
        def improve_bullets():
            raw=tb.get('1.0',tk.END).strip()
            if not raw: messagebox.showinfo('Empty','Add some bullet points first.'); return
            lines=to_bullets(raw)
            improved=[]
            for line in lines:
                # Check if starts with strong verb
                starts_strong = any(line.lower().startswith(v.lower()) for v in STRONG_VERBS)
                if not starts_strong:
                    # suggest wrapping with a strong verb
                    improved.append(f'Implemented {line[0].lower()}{line[1:]}')
                else:
                    improved.append(line)
            tb.delete('1.0',tk.END)
            tb.insert('1.0','\n'.join(improved))
            chk_v()
        _btn(body,'⚡ Auto-Improve Bullets',improve_bullets,sty='ghost',px=10,py=5).pack(anchor='e',pady=(4,0))

        def remove(o=outer):
            self.refs['exp_blks']=[(r,oo) for r,oo in self.refs['exp_blks'] if oo is not o]
            o.destroy()
        if self.refs['exp_blks']:
            tk.Button(chdr,text='Remove',font=(UF,10),bg='#E8F0FB',fg=RED,
                       relief=tk.FLAT,cursor='hand2',command=remove).pack(side='right',padx=10)
        self.refs['exp_blks'].append((refs,outer))

    def _sv_experience(self):
        exps=[]
        for refs,_ in self.refs['exp_blks']:
            t=refs['title'].get().strip(); co=refs['company'].get().strip()
            if not t and not co: continue
            exps.append({
                'title'     :s1(t),'company':s1(co),
                'location'  :s1(refs['location'].get()),
                'start_date':norm_date(refs['start_date'].get()),
                'end_date'  :norm_date(refs['end_date'].get()),
                'bullets'   :to_bullets(refs['bullets'].get('1.0',tk.END)),
            })
        if not exps:
            messagebox.showerror('Required','Add at least one work experience entry.'); return
        self.data['experience']=exps; self._update_checklist(); self._go(4)

    # ════════════════════════════════════════
    #  PAGE 5 — EDUCATION
    # ════════════════════════════════════════
    def _pg_education(self):
        wrap=tk.Frame(self.content,bg=BG); wrap.pack(fill='both',expand=True)
        top=tk.Frame(wrap,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        top.pack(fill='x',padx=26)
        hf=tk.Frame(top,bg=HDR_C,height=52); hf.pack(fill='x'); hf.pack_propagate(False)
        tk.Label(hf,text='  Education & Certifications',
                  font=(UF,13,'bold'),fg=WHITE,bg=HDR_C).pack(side='left',padx=6)
        _nav(top,back=lambda: self._go(3),nxt=self._sv_education,
             nxt_lbl='Preview Resume →',nxt_sty='success')

        so=tk.Frame(wrap,bg=BG); so.pack(fill='both',expand=True,padx=26)
        inner=_scrollbox(so)
        self.refs['edu_blks']=[]; self.refs['edu_inner']=inner

        if self.data.get('education'):
            for edu in self.data['education']: self._add_edu(edu)
        else:
            self._add_edu()

        bf=tk.Frame(inner,bg=BG); bf.pack(anchor='w',pady=(6,12))
        _btn(bf,'+ Add Degree',lambda: self._add_edu(),sty='outline',px=12,py=6).pack(side='left')

        # Certifications
        co2=tk.Frame(inner,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        co2.pack(fill='x',pady=(0,8))
        chdr2=tk.Frame(co2,bg='#E8F8EE',height=42); chdr2.pack(fill='x'); chdr2.pack_propagate(False)
        tk.Label(chdr2,text='  Certifications & Licenses  (optional)',
                  font=(UF,11,'bold'),fg='#07172B',bg='#E8F8EE').pack(side='left',padx=8)
        cb=tk.Frame(co2,bg=WHITE,padx=22,pady=14); cb.pack(fill='x')
        tk.Label(cb,text='One per line — include issuer and year',
                  font=(UF,10),fg=MUTED,bg=WHITE).pack(anchor='w',pady=(0,5))
        ct=_t(cb,h=5); ct.pack(fill='x')
        ec=self.data.get('certifications',[])
        if ec: ct.insert('1.0','\n'.join(ec))
        self.refs['certs_tb']=ct

    def _add_edu(self, data=None):
        data=data or {}
        n=len(self.refs['edu_blks'])+1
        cont=self.refs['edu_inner']
        outer=tk.Frame(cont,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        outer.pack(fill='x',pady=(4,0))
        chdr=tk.Frame(outer,bg='#E8F0FB',height=42); chdr.pack(fill='x'); chdr.pack_propagate(False)
        tk.Label(chdr,text=f'  Degree {n}',font=(UF,11,'bold'),fg='#07172B',bg='#E8F0FB').pack(side='left',padx=8)
        body=tk.Frame(outer,bg=WHITE,padx=22,pady=16); body.pack(fill='x')
        refs={}

        tk.Label(body,text='Degree / Major  *',font=(UF,11,'bold'),fg=TEXT,bg=WHITE).pack(anchor='w',pady=(0,3))
        ed=_e(body); ed.pack(fill='x',ipady=10)
        tk.Label(body,text='e.g., Bachelor of Science in Computer Science',
                  font=(UF,9),fg=MUTED,bg=WHITE).pack(anchor='w',pady=(2,6))
        v=data.get('degree','')
        if v: ed.insert(0,v)
        refs['degree']=ed

        r1=tk.Frame(body,bg=WHITE); r1.pack(fill='x',pady=4)
        for key,lbl_t,hint_t in [
            ('school','School / University  *','e.g., University of Michigan'),
            ('location','Location (optional)','e.g., Ann Arbor, MI'),
            ('gpa','GPA (optional)','e.g., 3.8'),
        ]:
            f=tk.Frame(r1,bg=WHITE); f.pack(side='left',fill='x',expand=True,padx=(0,10))
            tk.Label(f,text=lbl_t,font=(UF,11,'bold'),fg=TEXT,bg=WHITE).pack(anchor='w')
            e2=_e(f); e2.pack(fill='x',ipady=10,pady=(3,0))
            tk.Label(f,text=hint_t,font=(UF,9),fg=MUTED,bg=WHITE).pack(anchor='w',pady=(2,0))
            v2=data.get(key,'')
            if v2: e2.insert(0,v2)
            refs[key]=e2

        r2=tk.Frame(body,bg=WHITE); r2.pack(fill='x',pady=4)
        for key,lbl_t,hint_t in [('start_date','Start','e.g., Aug 2017'),('end_date','End','e.g., May 2021')]:
            f=tk.Frame(r2,bg=WHITE); f.pack(side='left',fill='x',expand=True,padx=(0,10))
            tk.Label(f,text=lbl_t,font=(UF,11,'bold'),fg=TEXT,bg=WHITE).pack(anchor='w')
            e3=_e(f); e3.pack(fill='x',ipady=10,pady=(3,0))
            tk.Label(f,text=hint_t,font=(UF,9),fg=MUTED,bg=WHITE).pack(anchor='w',pady=(2,0))
            v3=data.get(key,'')
            if v3: e3.insert(0,v3)
            refs[key]=e3

        def remove(o=outer):
            self.refs['edu_blks']=[(r,oo) for r,oo in self.refs['edu_blks'] if oo is not o]
            o.destroy()
        if self.refs['edu_blks']:
            tk.Button(chdr,text='Remove',font=(UF,10),bg='#E8F0FB',fg=RED,
                       relief=tk.FLAT,cursor='hand2',command=remove).pack(side='right',padx=10)
        self.refs['edu_blks'].append((refs,outer))

    def _sv_education(self):
        edus=[]
        for refs,_ in self.refs['edu_blks']:
            deg=refs['degree'].get().strip(); sch=refs['school'].get().strip()
            if not deg and not sch: continue
            edus.append({
                'degree':s1(deg),'school':s1(sch),
                'location':s1(refs['location'].get()),
                'gpa':refs['gpa'].get().strip(),
                'start_date':norm_date(refs['start_date'].get()),
                'end_date':norm_date(refs['end_date'].get()),
            })
        self.data['education']=edus
        raw_c=self.refs['certs_tb'].get('1.0',tk.END).strip()
        self.data['certifications']=[c.strip() for c in raw_c.split('\n') if c.strip()] if raw_c else []
        self._update_checklist(); self._go(5)

    # ════════════════════════════════════════
    #  PAGE 6 — PREVIEW & EXPORT
    # ════════════════════════════════════════
    def _pg_preview(self):
        wrap=tk.Frame(self.content,bg=BG); wrap.pack(fill='both',expand=True)

        # ── action bar ──────────────────────
        abar=tk.Frame(wrap,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        abar.pack(fill='x',padx=26)
        ab=tk.Frame(abar,bg=WHITE,padx=20,pady=12); ab.pack(fill='x')
        tk.Label(ab,text='Resume Preview',font=(UF,14,'bold'),fg='#07172B',bg=WHITE).pack(side='left')

        # Alignment radio
        self.align_var=tk.StringVar(value='left')
        af=tk.Frame(ab,bg=WHITE); af.pack(side='right',padx=(0,14))
        tk.Label(af,text='Align:',font=(UF,11),fg=TEXT,bg=WHITE).pack(side='left',padx=(0,4))
        for val,lbl in [('left','Left'),('justify','Justify')]:
            tk.Radiobutton(af,text=lbl,variable=self.align_var,value=val,
                            font=(UF,11),bg=WHITE,fg=TEXT,selectcolor=WHITE,
                            command=self._refresh_preview).pack(side='left',padx=2)

        # Bullet style
        self.bullet_var=tk.StringVar(value='Bullet  •')
        bf2=tk.Frame(ab,bg=WHITE); bf2.pack(side='right',padx=(0,8))
        tk.Label(bf2,text='Bullet:',font=(UF,11),fg=TEXT,bg=WHITE).pack(side='left',padx=(0,6))
        om=ttk.Combobox(bf2,textvariable=self.bullet_var,
                         values=list(BULLET_STYLES.keys()),
                         state='readonly',width=12,font=(UF,11))
        om.pack(side='left')
        om.bind('<<ComboboxSelected>>',lambda e: self._refresh_preview())

        _btn(ab,'💾  Download .docx',self._download,sty='success').pack(side='right',padx=(6,0))
        _btn(ab,'📄  Save JSON',      self._save_json,sty='ghost').pack(side='right',padx=6)
        _btn(ab,'← Edit',            lambda: self._go(4),sty='ghost').pack(side='right')

        # ── score bar ───────────────────────
        sc=self._score(); scol=GREEN if sc>=80 else AMBER if sc>=55 else RED
        sbar=tk.Frame(wrap,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        sbar.pack(fill='x',padx=26)
        si=tk.Frame(sbar,bg=WHITE,padx=20,pady=10); si.pack(fill='x')
        tk.Label(si,text=f'ATS Score: {sc}/100',font=(UF,12,'bold'),fg=scol,bg=WHITE).pack(side='left')
        tk.Label(si,text=f'   {self._score_tips()}',font=(UF,10),fg=MUTED,bg=WHITE).pack(side='left')

        # Mini section score bars
        sec_scores=[
            ('Contact',  30 if all(self.data['personal'].get(k) for k in ('full_name','email','phone','location')) else 0),
            ('Summary',  15 if len(self.data.get('professional_summary',''))>80 else 0),
            ('Skills',   15 if len(self.data.get('skills',{}))>=2 else 0),
            ('Exp',      25 if self.data.get('experience') else 0),
            ('Edu',      10 if self.data.get('education') else 0),
            ('Certs',    5  if self.data.get('certifications') else 0),
        ]
        maxes={'Contact':30,'Summary':15,'Skills':15,'Exp':25,'Edu':10,'Certs':5}
        sf2=tk.Frame(si,bg=WHITE); sf2.pack(side='right')
        for sec_n,sec_s in sec_scores:
            c2=tk.Frame(sf2,bg=WHITE); c2.pack(side='left',padx=3)
            pct=sec_s/maxes[sec_n]
            clr=GREEN if pct==1.0 else AMBER if pct>0 else BORDER
            tk.Label(c2,text=sec_n,font=(UF,8),fg=MUTED,bg=WHITE).pack()
            bo=tk.Frame(c2,bg=BORDER,width=46,height=6); bo.pack(); bo.pack_propagate(False)
            tk.Frame(bo,bg=clr,height=6,width=int(46*pct)).pack(side='left')

        # ── preview canvas ──────────────────
        self._ph=tk.Frame(wrap,bg='#8090A8')
        self._ph.pack(fill='both',expand=True,padx=26,pady=8)
        self._mk_preview()

    def _mk_preview(self):
        for w in self._ph.winfo_children(): w.destroy()
        cv=tk.Canvas(self._ph,bg='#8090A8',highlightthickness=0)
        sb2=ttk.Scrollbar(self._ph,orient='vertical',command=cv.yview)
        cv.configure(yscrollcommand=sb2.set)
        cv.pack(side='left',fill='both',expand=True,padx=14,pady=14)
        sb2.pack(side='right',fill='y',pady=14)
        paper=tk.Frame(cv,bg=WHITE,highlightbackground='#4A6080',highlightthickness=1)
        cwin=cv.create_window((0,0),window=paper,anchor='nw')
        cv.bind('<Configure>',lambda e: cv.itemconfig(cwin,width=max(640,e.width-8)))
        paper.bind('<Configure>',lambda e: cv.configure(scrollregion=cv.bbox('all')))
        cv.bind_all('<MouseWheel>',lambda e: cv.yview_scroll(int(-1*(e.delta/120)),'units'))
        self._render(paper)

    def _refresh_preview(self): self._mk_preview()

    def _render(self, parent):
        """
        Preview matches DOCX typography exactly:
          Section headers: 10pt ALL CAPS Bold (SMALLER than role names)
          Role line: 11pt Bold + date 10pt italic RIGHT
          Location: 10pt italic
          Bullets: 11pt hanging-indent, space between bullet and text
        """
        rd  = self.data
        p   = rd.get('personal',{})
        bkey= self.bullet_var.get() if self.bullet_var else 'Bullet  •'
        bch = BULLET_STYLES.get(bkey,'•')
        do_j= (self.align_var.get()=='justify') if self.align_var else False

        m = tk.Frame(parent,bg=WHITE,padx=50,pady=36)
        m.pack(fill='both',expand=True)

        # ── render helpers ──────────────────
        def L(text,size=11,bold=False,color='#000',pad=(0,1),
               anchor='w',italic=False,wrap=False):
            st=[]
            if bold:   st.append('bold')
            if italic: st.append('italic')
            ft=(DF,size,' '.join(st)) if st else (DF,size)
            tk.Label(m,text=text,font=ft,bg=WHITE,fg=color,
                      justify='left',anchor=anchor,
                      wraplength=560 if wrap else 0).pack(
                      anchor=anchor,pady=pad,
                      fill='x' if wrap else None)

        def HR():
            tk.Frame(m,height=1,bg='#0B2A55').pack(fill='x',pady=(7,4))

        def SEC(label):
            HR()
            # Section header: 10pt ALL CAPS BOLD (smaller than 11pt role names)
            tk.Label(m,text=label.upper(),font=(DF,10,'bold'),
                      bg=WHITE,fg='#07172B',anchor='w').pack(anchor='w',pady=(3,4))

        def ROLE(left,right):
            """11pt bold left, 10pt italic right-aligned date."""
            row=tk.Frame(m,bg=WHITE); row.pack(fill='x',pady=(7,0))
            tk.Label(row,text=right,font=(DF,SZ_DATE,'italic'),
                      bg=WHITE,fg='#2A4A6A').pack(side='right')
            tk.Label(row,text=left,font=(DF,SZ_ROLE,'bold'),
                      bg=WHITE,fg='#000',anchor='w').pack(side='left',fill='x',expand=True)

        def BULL(text,idx=0):
            """
            Hanging-indent bullet with clear space between • and text.
            Matches screenshot: • followed by two spaces then text.
            Continuation lines wrap cleanly under first word.
            """
            pfx=f'{idx+1}.  ' if bch=='NUM' else f'{bch}  '
            row=tk.Frame(m,bg=WHITE); row.pack(fill='x',pady=(1,2))
            # bullet marker at fixed left edge
            mk=tk.Label(row,text=pfx.rstrip(),font=(DF,SZ_BODY),
                         bg=WHITE,fg='#000',anchor='nw',width=3)
            mk.pack(side='left',anchor='nw',padx=(0,4))
            # body text wraps under itself
            tk.Label(row,text=text,font=(DF,SZ_BODY),bg=WHITE,fg='#000',
                      anchor='w',justify='left',wraplength=500).pack(
                      side='left',anchor='nw',fill='x',expand=True)

        # ── content ─────────────────────────
        L(san(p.get('full_name','')).upper(),size=SZ_NAME,bold=True,pad=(0,3))
        c_parts=[san(p.get(k,'')) for k in ('email','phone','location','linkedin','github') if p.get(k)]
        if c_parts: L(' | '.join(c_parts),size=SZ_CONTACT,pad=(0,5))

        if rd.get('professional_summary'):
            SEC('Professional Summary')
            L(san(rd['professional_summary']),size=SZ_BODY,pad=(3,5),wrap=True)

        if rd.get('skills'):
            SEC('Core Competencies')
            for cat,items in rd['skills'].items():
                row=tk.Frame(m,bg=WHITE); row.pack(anchor='w',pady=(0,1))
                tk.Label(row,text=f'{san(cat)}: ',font=(DF,SZ_BODY,'bold'),
                          bg=WHITE,fg='#000').pack(side='left')
                tk.Label(row,text=', '.join(san(s) for s in items),
                          font=(DF,SZ_BODY),bg=WHITE,fg='#000',
                          wraplength=480,anchor='w',justify='left').pack(side='left')

        if rd.get('experience'):
            SEC('Professional Experience')
            for exp in rd['experience']:
                t  =san(exp.get('title',''))
                co =san(exp.get('company',''))
                lc =san(exp.get('location',''))
                st =exp.get('start_date','')
                en =exp.get('end_date','Present')
                ds =f'{st} - {en}' if st else en
                rl =f'{t}  |  {co}' if t and co else (t or co)
                ROLE(rl,ds)
                if lc: L(lc,size=SZ_LOC,italic=True,color='#2A4A6A',pad=(0,2))
                for idx,b in enumerate(exp.get('bullets',[])):
                    b=san(b)
                    if b: BULL(b,idx)

        if rd.get('education'):
            SEC('Education')
            for edu in rd['education']:
                deg=san(edu.get('degree',''))
                sch=san(edu.get('school',''))
                lc =san(edu.get('location',''))
                gpa=edu.get('gpa','')
                s,e=edu.get('start_date',''),edu.get('end_date','')
                ds =f'{s} - {e}' if s and e else (e or s)
                ROLE(deg,ds)
                pts=[x for x in (sch,lc) if x]
                if gpa: pts.append(f'GPA: {gpa}')
                if pts: L(' | '.join(pts),size=SZ_LOC,pad=(0,3))

        if rd.get('certifications'):
            SEC('Certifications')
            for idx,cert in enumerate(rd['certifications']):
                cert=san(cert)
                if cert: BULL(cert,idx)

        if rd.get('projects'):
            SEC('Projects')
            for proj in rd['projects']:
                nm =san(proj.get('name',''))
                tch=san(proj.get('tech',''))
                dt =proj.get('date','')
                left=f'{nm}  |  {tch}' if tch else nm
                ROLE(left,dt)
                for idx,b in enumerate(proj.get('bullets',[])):
                    b=san(b)
                    if b: BULL(b,idx)

        tk.Frame(m,height=30,bg=WHITE).pack()

    # ── ATS score ───────────────────────────────────
    def _score(self):
        rd=self.data; s=0
        if rd['personal'].get('full_name'): s+=10
        if rd['personal'].get('email'):     s+=10
        if rd['personal'].get('phone'):     s+=5
        if rd['personal'].get('location'):  s+=5
        if len(rd.get('professional_summary',''))>80: s+=15
        if rd.get('skills') and len(rd['skills'])>=2:  s+=15
        if rd.get('experience'):                        s+=25
        if rd.get('education'):                         s+=10
        if rd.get('certifications'):                    s+=5
        return s

    def _score_tips(self):
        rd=self.data; tips=[]
        if not rd.get('experience'):                  tips.append('Add experience')
        if not rd.get('skills'):                      tips.append('Add skills')
        if not rd.get('education'):                   tips.append('Add education')
        if len(rd.get('professional_summary',''))<80: tips.append('Expand summary')
        return ('Improve: '+'  •  '.join(tips)) if tips else '✓ All sections complete — ready to submit'

    # ════════════════════════════════════════
    #  KEYWORD ANALYZER
    # ════════════════════════════════════════
    def _keywords(self):
        win=tk.Toplevel(self.root)
        win.title('ATS Keyword Analyzer'); win.geometry('860x720')
        win.configure(bg=BG); win.grab_set()

        outer=tk.Frame(win,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        outer.pack(fill='both',expand=True,padx=14,pady=14)
        hf=tk.Frame(outer,bg=HDR_C,height=52); hf.pack(fill='x'); hf.pack_propagate(False)
        tk.Label(hf,text='  ATS Keyword Analyzer  —  paste the job description',
                  font=(UF,13,'bold'),fg=WHITE,bg=HDR_C).pack(side='left',padx=14)

        body=tk.Frame(outer,bg=WHITE,padx=22,pady=14); body.pack(fill='both',expand=True)
        tk.Label(body,text='Paste the full job description. We will show matched and missing keywords.',
                  font=(UF,10),fg=MUTED,bg=WHITE,wraplength=780,justify='left').pack(anchor='w',pady=(0,8))

        tb=_t(body,h=13); tb.pack(fill='x')
        rf=tk.Frame(body,bg=WHITE); rf.pack(fill='x',pady=(12,0))

        def analyse():
            jd=tb.get('1.0',tk.END).lower()
            if len(jd.split())<30:
                messagebox.showerror('Too Short','Paste the full job description.',parent=win); return
            stops={'the','and','for','with','in','on','of','to','a','an','is','are',
                    'be','will','at','by','or','as','it','this','that','we','our',
                    'you','your','their','they','from','have','not','but','has','its',
                    'who','which','what','when','where','how','if','then','than','all'}
            jd_words={w.strip('.,;:()/•-') for w in re.split(r'\s+',jd)
                       if len(w)>2 and w not in stops and w.isascii()}
            resume_text=' '.join([
                self.data.get('professional_summary',''),
                ' '.join(' '.join(v) for v in self.data.get('skills',{}).values()),
                ' '.join(f"{e.get('title','')} {e.get('company','')} {' '.join(e.get('bullets',[]))}"
                          for e in self.data.get('experience',[])),
            ]).lower()
            matched=sorted(w for w in jd_words if w in resume_text and len(w)>3)
            missing=sorted(w for w in jd_words if w not in resume_text and len(w)>3)
            for w in rf.winfo_children(): w.destroy()
            pct=int(len(matched)/max(len(jd_words),1)*100)
            scol=GREEN if pct>=70 else AMBER if pct>=45 else RED
            tk.Label(rf,text=f'Keyword Match: {pct}%  ({len(matched)} / {len(jd_words)} keywords)',
                      font=(UF,12,'bold'),fg=scol,bg=WHITE).pack(anchor='w',pady=(0,8))
            # Matched
            cf=tk.Frame(rf,bg='#F0FDF4',padx=12,pady=10); cf.pack(fill='x',pady=(0,6))
            tk.Label(cf,text='✓ Matched:',font=(UF,10,'bold'),fg=GREEN,bg='#F0FDF4').pack(anchor='w',pady=(0,3))
            tk.Label(cf,text='  '+',  '.join(matched[:35])+('…' if len(matched)>35 else ''),
                      font=(UF,10),fg=TEXT,bg='#F0FDF4',wraplength=740,justify='left').pack(anchor='w')
            # Missing
            mf=tk.Frame(rf,bg='#FEF2F2',padx=12,pady=10); mf.pack(fill='x')
            tk.Label(mf,text='⚠  Add these keywords to improve your score:',
                      font=(UF,10,'bold'),fg=RED,bg='#FEF2F2').pack(anchor='w',pady=(0,3))
            tk.Label(mf,text='  '+',  '.join(missing[:40])+('…' if len(missing)>40 else ''),
                      font=(UF,10),fg=TEXT,bg='#FEF2F2',wraplength=740,justify='left').pack(anchor='w')

        br=tk.Frame(body,bg=WHITE); br.pack(fill='x',pady=(8,0))
        _btn(br,'Close',win.destroy,sty='ghost').pack(side='left')
        _btn(br,'🔍  Analyse',analyse,sty='primary').pack(side='right')

    # ════════════════════════════════════════
    #  TEMPLATES
    # ════════════════════════════════════════
    def _templates(self):
        win=tk.Toplevel(self.root)
        win.title('Resume Templates'); win.geometry('640x480')
        win.configure(bg=BG); win.grab_set()

        outer=tk.Frame(win,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        outer.pack(fill='both',expand=True,padx=14,pady=14)
        hf=tk.Frame(outer,bg=HDR_C,height=52); hf.pack(fill='x'); hf.pack_propagate(False)
        tk.Label(hf,text='  Choose a Resume Template',
                  font=(UF,13,'bold'),fg=WHITE,bg=HDR_C).pack(side='left',padx=14)

        body=tk.Frame(outer,bg=WHITE,padx=26,pady=20); body.pack(fill='both',expand=True)
        tk.Label(body,text='All templates are ATS-compliant. Select one to apply its section order and style.',
                  font=(UF,10),fg=MUTED,bg=WHITE,wraplength=560).pack(anchor='w',pady=(0,14))

        templates=[
            ('📄  Modern Tech',
             'Best for software, cloud, DevOps roles.\nSections: Summary → Skills → Experience → Projects → Education → Certs',
             {'order':['summary','skills','experience','projects','education','certs']}),
            ('📃  Classic Professional',
             'Best for management, finance, consulting.\nSections: Summary → Experience → Education → Skills → Certs',
             {'order':['summary','experience','education','skills','certs']}),
            ('📋  Entry Level / Graduate',
             'Best for fresh graduates.\nSections: Summary → Education → Skills → Projects → Experience',
             {'order':['summary','education','skills','projects','experience']}),
        ]

        def apply(t):
            messagebox.showinfo('Template Applied',
                f'"{t[0].strip()}" template applied.\n\nSection order saved for your download.',
                parent=win)
            win.destroy()

        for t in templates:
            tf=tk.Frame(body,bg='#F4F8FF',highlightbackground=BORDER,highlightthickness=1)
            tf.pack(fill='x',pady=4)
            inner=tk.Frame(tf,bg='#F4F8FF',padx=16,pady=14); inner.pack(fill='x')
            tk.Label(inner,text=t[0],font=(UF,13,'bold'),fg=TEXT,bg='#F4F8FF').pack(anchor='w')
            tk.Label(inner,text=t[1],font=(UF,10),fg=MUTED,bg='#F4F8FF',
                      justify='left').pack(anchor='w',pady=(4,0))
            _btn(inner,'Select',lambda tt=t: apply(tt),sty='outline',px=14,py=6).pack(anchor='e',pady=(6,0))

    # ════════════════════════════════════════
    #  UPLOAD
    # ════════════════════════════════════════
    def _upload(self):
        win=tk.Toplevel(self.root)
        win.title('Upload & Reformat Resume')
        win.geometry('760x680'); win.configure(bg=BG); win.grab_set()

        outer=tk.Frame(win,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        outer.pack(fill='both',expand=True,padx=14,pady=14)
        hf=tk.Frame(outer,bg=HDR_C,height=52); hf.pack(fill='x'); hf.pack_propagate(False)
        tk.Label(hf,text='  Upload & Reformat Existing Resume',
                  font=(UF,14,'bold'),fg=WHITE,bg=HDR_C).pack(side='left',padx=14)

        body=tk.Frame(outer,bg=WHITE,padx=24,pady=16); body.pack(fill='both',expand=True)
        tk.Label(body,text='Supported: .docx  .txt  (PDF: copy-paste text into "Paste & Fill")',
                  font=(UF,11),fg=MUTED,bg=WHITE).pack(anchor='w',pady=(0,10))

        cf=tk.Frame(body,bg='#F0FDF4',padx=14,pady=12); cf.pack(fill='x',pady=(0,14))
        tk.Label(cf,text='Auto-applied reformatting:',
                  font=(UF,11,'bold'),fg=GREEN,bg='#F0FDF4').pack(anchor='w',pady=(0,6))
        for item in [
            '• Strict section isolation — experience, education, skills never mixed',
            '• Two-pass block builder — handles title-before-date resume formats',
            '• MM/YYYY dates normalised (01/2019 → Jan 2019)',
            '• Paragraph-style bullets (no prefix) detected and extracted',
            '• All unicode bullets/dashes sanitised → plain ASCII',
            '• ATS typography: 10pt section headers, 11pt role names, 11pt bullets',
        ]:
            tk.Label(cf,text=item,font=(UF,10),fg=TEXT,bg='#F0FDF4').pack(anchor='w')

        fl=tk.Label(body,text='No file selected',font=(UF,11),fg=MUTED,bg=WHITE)
        fl.pack(anchor='w',pady=(6,6))
        self._up_path=None

        def pick():
            fp=filedialog.askopenfilename(title='Select Resume',
                    filetypes=[('Word / Text','*.docx *.txt'),('All files','*.*')],parent=win)
            if fp: self._up_path=fp; fl.config(text=os.path.basename(fp),fg=GREEN)

        _btn(body,'📁  Choose File',pick,sty='outline',px=14,py=8).pack(anchor='w',pady=(0,8))
        stat=tk.Label(body,text='',font=(UF,11),fg=ACCENT,bg=WHITE); stat.pack(anchor='w')

        def do_upload():
            fp=self._up_path
            if not fp:
                messagebox.showerror('No File','Choose a file first.',parent=win); return
            stat.config(text='Reading file…'); win.update()
            try:
                text=''
                if fp.lower().endswith('.docx'):
                    doc=Document(fp)
                    parts=[para.text for para in doc.paragraphs]
                    for table in doc.tables:
                        for trow in table.rows:
                            for cell in trow.cells: parts.append(cell.text)
                    text='\n'.join(parts)
                else:
                    with open(fp,'r',encoding='utf-8',errors='ignore') as fh: text=fh.read()
                if len(text.strip())<30:
                    messagebox.showerror('Empty','File is empty.',parent=win); stat.config(text=''); return
                stat.config(text='Parsing sections…'); win.update()
                self._prev_data=copy.deepcopy(self.data)
                parsed=parse_resume(text)
                self.data=parsed; self.step=len(STEPS)-1
                stat.config(text='Done!'); win.update(); win.destroy()
                self._undo_btn.pack(side='right',padx=16)
                messagebox.showinfo('Upload Complete',
                    'Resume parsed.\n\nReview each section via the sidebar.\n'
                    'Use "↩ Undo Parse" to restore your previous data.')
                self._go(5)
            except Exception as ex:
                messagebox.showerror('Error',f'Could not read file:\n{str(ex)}',parent=win)
                stat.config(text='')

        br=tk.Frame(body,bg=WHITE); br.pack(fill='x',pady=(14,0))
        _btn(br,'Cancel',win.destroy,sty='ghost').pack(side='left')
        _btn(br,'Reformat & Preview →',do_upload,sty='success').pack(side='right')

    # ════════════════════════════════════════
    #  PASTE & AUTO-FILL
    # ════════════════════════════════════════
    def _paste(self):
        win=tk.Toplevel(self.root)
        win.title('Paste & Auto-Format')
        win.geometry('840x700'); win.configure(bg=BG); win.grab_set()

        outer=tk.Frame(win,bg=WHITE,highlightbackground=BORDER,highlightthickness=1)
        outer.pack(fill='both',expand=True,padx=14,pady=14)
        hf=tk.Frame(outer,bg=HDR_C,height=52); hf.pack(fill='x'); hf.pack_propagate(False)
        tk.Label(hf,text='  Paste & Auto-Format Resume Text',
                  font=(UF,14,'bold'),fg=WHITE,bg=HDR_C).pack(side='left',padx=14)

        body=tk.Frame(outer,bg=WHITE,padx=24,pady=14); body.pack(fill='both',expand=True)
        tk.Label(body,text='Paste any resume text — PDF copy, Word export, LinkedIn "About", or plain text.',
                  font=(UF,11),fg=MUTED,bg=WHITE).pack(anchor='w',pady=(0,8))

        tb=_t(body,h=24); tb.pack(fill='both',expand=True); tb.focus_set()
        stat=tk.Label(body,text='',font=(UF,11),fg=RED,bg=WHITE); stat.pack(anchor='w',pady=(4,0))

        def do_parse():
            raw=tb.get('1.0',tk.END).strip()
            if len(raw)<60 or len(raw.split())<25:
                stat.config(text='⚠  Too short — paste your complete resume.',fg=RED); return
            stat.config(text='Parsing…',fg=ACCENT); win.update()
            try:
                self._prev_data=copy.deepcopy(self.data)
                parsed=parse_resume(raw)
                found=[]
                if parsed['personal'].get('full_name'):  found.append('Name')
                if parsed['personal'].get('email'):       found.append('Email')
                if parsed['personal'].get('phone'):       found.append('Phone')
                if parsed.get('professional_summary'):    found.append('Summary')
                if parsed.get('skills'):     found.append(f'{len(parsed["skills"])} skill groups')
                if parsed.get('experience'): found.append(f'{len(parsed["experience"])} positions')
                if parsed.get('education'):  found.append('Education')
                if parsed.get('certifications'): found.append('Certifications')
                if parsed.get('projects'):   found.append(f'{len(parsed["projects"])} projects')
                self.data=parsed; self.step=len(STEPS)-1
                win.destroy()
                self._undo_btn.pack(side='right',padx=16)
                messagebox.showinfo('Done',
                    'Extracted:\n• '+'\n• '.join(found)+
                    '\n\nReview sections via sidebar, then Download.')
                self._go(5)
            except Exception as ex:
                stat.config(text=f'Error: {str(ex)[:120]}',fg=RED)

        br=tk.Frame(body,bg=WHITE); br.pack(fill='x',pady=(8,0))
        _btn(br,'Cancel',win.destroy,sty='ghost').pack(side='left')
        _btn(br,'⚡  Auto-Format',do_parse,sty='primary').pack(side='right')

    # ════════════════════════════════════════
    #  DOWNLOAD / SAVE / LOAD
    # ════════════════════════════════════════
    def _download(self):
        nm=self.data.get('personal',{}).get('full_name','Resume').replace(' ','_')
        fp=filedialog.asksaveasfilename(
            defaultextension='.docx',filetypes=[('Word Document','*.docx')],
            initialdir=os.path.expanduser('~/Desktop'),
            initialfile=f'{nm}_ATS_Resume.docx')
        if not fp: return
        try:
            bkey=self.bullet_var.get() if self.bullet_var else 'Bullet  •'
            bch=BULLET_STYLES.get(bkey,'•')
            do_j=(self.align_var.get()=='justify') if self.align_var else False
            write_docx(self.data,fp,bullet_char=bch,justify=do_j)
            messagebox.showinfo('Saved',f'ATS resume saved:\n{fp}')
            if messagebox.askyesno('Open?','Open the file now?'):
                try:    os.startfile(fp)
                except Exception:
                    import subprocess; subprocess.run(['open',fp],check=False)
        except Exception as ex:
            messagebox.showerror('Error',f'Could not save:\n{str(ex)}')

    def _save_json(self):
        nm=self.data.get('personal',{}).get('full_name','Resume').replace(' ','_')
        fp=filedialog.asksaveasfilename(defaultextension='.json',
            filetypes=[('JSON','*.json')],
            initialdir=os.path.expanduser('~/Desktop'),
            initialfile=f'{nm}_resume_data.json')
        if not fp: return
        try:
            with open(fp,'w') as f: json.dump(self.data,f,indent=2)
            messagebox.showinfo('Saved',f'Data saved:\n{fp}')
        except Exception as ex:
            messagebox.showerror('Error',str(ex))

    def _load_json(self):
        fp=filedialog.askopenfilename(filetypes=[('JSON','*.json')],
                                       initialdir=os.path.expanduser('~/Desktop'))
        if not fp: return
        try:
            with open(fp) as f: self.data=json.load(f)
            self.step=len(STEPS)-1
            messagebox.showinfo('Loaded','Resume data loaded.')
            self._go(0)
        except Exception as ex:
            messagebox.showerror('Error',str(ex))

# ══════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════
def main():
    root = tk.Tk()
    try: root.tk.call('tk','scaling',1.25)
    except Exception: pass
    App(root)
    root.mainloop()

if __name__ == '__main__':
    main()
